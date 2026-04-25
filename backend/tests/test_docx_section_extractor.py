"""Tests for services.docx_section_extractor."""
from __future__ import annotations

import io

from docx import Document

from services.docx_section_extractor import (
    TableBlock,
    TextBlock,
    extract_sections,
)


def _make_docx(builder) -> bytes:
    doc = Document()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extracts_heading_styled_sections():
    def build(doc):
        h = doc.add_paragraph("1. Date despre program")
        h.style = doc.styles["Heading 1"]
        doc.add_paragraph("Universitatea Transilvania")
        h2 = doc.add_paragraph("2. Date despre disciplina")
        h2.style = doc.styles["Heading 1"]
        doc.add_paragraph("Analiza matematica")

    sections = extract_sections(_make_docx(build))

    headings = [s.heading for s in sections]
    assert "1. Date despre program" in headings
    assert "2. Date despre disciplina" in headings
    sec1 = next(s for s in sections if s.heading == "1. Date despre program")
    assert isinstance(sec1.body[0], TextBlock)
    assert sec1.body[0].paragraphs == ["Universitatea Transilvania"]


def test_extracts_numbering_pattern_when_no_heading_style():
    def build(doc):
        doc.add_paragraph("8.1 Tematica activitatilor de curs")
        doc.add_paragraph("Curs 1: Limite de functii")

    sections = extract_sections(_make_docx(build))

    assert any(s.heading.startswith("8.1") for s in sections)
    sec = next(s for s in sections if s.heading.startswith("8.1"))
    assert sec.body[0].paragraphs == ["Curs 1: Limite de functii"]


def test_includes_tables_in_body():
    def build(doc):
        h = doc.add_paragraph("3. Timpul total estimat")
        h.style = doc.styles["Heading 1"]
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Curs"
        table.rows[0].cells[1].text = "28"
        table.rows[1].cells[0].text = "Seminar"
        table.rows[1].cells[1].text = "14"

    sections = extract_sections(_make_docx(build))

    sec = next(s for s in sections if s.heading.startswith("3."))
    table_blocks = [b for b in sec.body if isinstance(b, TableBlock)]
    assert len(table_blocks) == 1
    # Either headers+rows split or full rows preserved.
    all_cells = [c for r in table_blocks[0].rows for c in r] + table_blocks[0].headers
    assert "Curs" in all_cells and "28" in all_cells
    assert "Seminar" in all_cells and "14" in all_cells


def test_empty_doc_returns_single_preamble():
    def build(doc):
        doc.add_paragraph("Just a note")

    sections = extract_sections(_make_docx(build))

    assert len(sections) == 1
    assert sections[0].level == 0
    assert sections[0].body[0].paragraphs == ["Just a note"]


def test_normalised_heading_strips_diacritics_and_numbering():
    def build(doc):
        h = doc.add_paragraph("8.1 Tematica activităților de curs")
        h.style = doc.styles["Heading 2"]

    sec = next(
        s for s in extract_sections(_make_docx(build))
        if s.heading.startswith("8.1")
    )
    assert sec.heading_norm == "tematica activitatilor de curs"
